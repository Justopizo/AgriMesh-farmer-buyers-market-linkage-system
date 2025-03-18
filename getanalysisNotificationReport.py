import os
import psycopg2
from docx import Document
from PyQt6.QtWidgets import QTextEdit
from adminDashboard import Ui_admnidashboardDialog

def generate_report():
    try:
        conn = psycopg2.connect(
            host="localhost",
            user="postgres",
            password="3062",
            database="agrimesh"
        )
        cursor = conn.cursor()
        
        doc = Document()
        report_text = ""

        # Sales Data Report
        doc.add_heading('Sales Data Report', level=1)
        report_text += "Sales Data Report\n\n"
        cursor.execute("SELECT buyername, phoneno, productname, quantity, price, status FROM vieworders;")
        sales_records = cursor.fetchall()
       
        sales_headers = ["Buyer Name", "Phone No", "Product Name", "Quantity", "Price", "Status"]
        report_text += "\t".join(sales_headers) + "\n"
        for row in sales_records:
            report_text += "\t".join(map(str, row)) + "\n"

        doc.add_paragraph(report_text)
        report_text += "\n\n"

        # Users Data Report
        doc.add_heading('Users Data Report', level=1)
        report_text += "Users Data Report\n\n"
        cursor.execute("SELECT name, password, role FROM userinfo;")
        users_records = cursor.fetchall()
        users_headers = ["User Name", "Password", "Role"]
        report_text += "\t".join(users_headers) + "\n"
        for row in users_records:
            report_text += "\t".join(map(str, row)) + "\n"

        doc.add_paragraph(report_text)
        
        cursor.close()
        conn.close()
        
        file_path = os.path.abspath("report.docx")
        doc.save(file_path)
        
        return report_text
    except Exception as e:
        return f"Error: {e}"

def getnoOfNewnotifications():
    conn = psycopg2.connect(
            host="localhost",
            user="postgres",
            password="3062",
            database="agrimesh"
        )
    cursor = conn.cursor()
    cursor.execute("SELECT buyername, phoneno, productname, quantity, price, status FROM vieworders;")
    sales_records = cursor.fetchall()
    return len(sales_records)

