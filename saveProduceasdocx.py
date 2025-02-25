import os
import psycopg2
from docx import Document
from PyQt6.QtWidgets import *

def save_to_docx(file_name):
    try:
        # Connect to PostgreSQL
        conn = psycopg2.connect(
                host="localhost",
                user="postgres",
                password="3062",
                database="agrimesh"
        )
        cursor = conn.cursor()

        # Fetch data
        query = "SELECT productname, category, quantity, price, location FROM produce"
        cursor.execute(query)
        records = cursor.fetchall()
        
        # Close connection
        cursor.close()
        conn.close()

        # Create a Word document
        doc = Document()
        doc.add_heading('Produce Data Report', level=1)

        # Add table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Product Name", "Category", "Quantity", "Price", "Location"]

        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Populate table with data
        for row_data in records:
            row_cells = table.add_row().cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)

        # Save the document
        file_path = os.path.abspath(file_name)
        doc.save(file_path)
        

        return file_path  

    except Exception as e:
        print("Error:", e)
        return None
