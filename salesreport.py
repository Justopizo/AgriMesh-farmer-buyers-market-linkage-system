import os
import psycopg2
from docx import Document
from PyQt6.QtWidgets import *

def save_to_docx(file_name):
    try:
        
        conn = psycopg2.connect(
                host="localhost",
                user="postgres",
                password="3062",
                database="agrimesh"
        )
        cursor = conn.cursor()

        
        query = "SELECT buyername, phoneno, productname, quantity, price, status FROM vieworders;"
        cursor.execute(query)
        records = cursor.fetchall()
        
        
        cursor.close()
        conn.close()

        
        doc = Document()
        doc.add_heading('Sales Data Report', level=1)

        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["Buyer Name", "phoneno", "productname","quantity","price","status"]

        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        
        for row_data in records:
            row_cells = table.add_row().cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data)

        
        file_path = os.path.abspath(file_name)
        doc.save(file_path)
        

        return file_path  

    except Exception as e:
        print("Error:", e)
        return None
